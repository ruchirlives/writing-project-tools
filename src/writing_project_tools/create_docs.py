from __future__ import annotations

import csv
import argparse
from datetime import datetime, timezone
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Pt


DEFAULT_AUTHOR = "Author"


def add_markdown_to_document(document: Document, markdown: str) -> None:
    for raw_line in markdown.splitlines():
        line = raw_line.strip()

        if not line:
            continue

        if line.startswith("# "):
            document.add_heading(line.removeprefix("# "), level=1)
        elif line.startswith("## "):
            document.add_heading(line.removeprefix("## "), level=2)
        elif line.startswith("### "):
            document.add_heading(line.removeprefix("### "), level=3)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:], style="List Bullet")
        elif ". " in line[:5] and line.split(". ", 1)[0].isdigit():
            document.add_paragraph(line.split(". ", 1)[1], style="List Number")
        else:
            document.add_paragraph(line)


def style_document(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Aptos Display"


def set_document_properties(
    document: Document,
    title: str,
    subject: str,
    author: str,
) -> None:
    timestamp = datetime.now(timezone.utc).replace(tzinfo=None)
    properties = document.core_properties
    properties.author = author
    properties.last_modified_by = author
    properties.title = title
    properties.subject = subject
    properties.keywords = ""
    properties.comments = ""
    properties.category = ""
    properties.content_status = ""
    properties.created = timestamp
    properties.modified = timestamp
    properties.revision = 1


def patch_extended_properties(docx_path: Path) -> None:
    with TemporaryDirectory(dir=docx_path.parent) as temp_dir:
        temp_path = Path(temp_dir) / docx_path.name

        with ZipFile(docx_path, "r") as source_zip, ZipFile(
            temp_path, "w", ZIP_DEFLATED
        ) as target_zip:
            for item in source_zip.infolist():
                data = source_zip.read(item.filename)

                if item.filename == "docProps/app.xml":
                    text = data.decode("utf-8")
                    text = text.replace(
                        "<Application>Microsoft Macintosh Word</Application>",
                        "<Application>Microsoft Word</Application>",
                    )
                    text = text.replace(
                        "<AppVersion>14.0000</AppVersion>",
                        "<AppVersion>16.0000</AppVersion>",
                    )
                    data = text.encode("utf-8")

                target_zip.writestr(item, data)

        temp_path.replace(docx_path)


def build_docx(
    source_path: Path,
    title: str,
    subject: str,
    output_path: Path,
    author: str,
) -> None:
    document = Document()
    style_document(document)
    set_document_properties(document, title, subject, author)

    add_markdown_to_document(document, source_path.read_text(encoding="utf-8"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    patch_extended_properties(output_path)


def read_config(config_path: Path) -> list[dict[str, str]]:
    if not config_path.exists():
        raise FileNotFoundError(
            f"Missing {config_path}. Create it with columns: "
            "source,title,subject,output,author"
        )

    with config_path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def generate_docs(project_dir: Path, config: str = "docx_sources.csv") -> list[Path]:
    root = project_dir.resolve()
    config_path = (root / config).resolve()
    outputs: list[Path] = []

    for row in read_config(config_path):
        source = row.get("source", "").strip()
        if not source:
            continue

        source_path = root / source
        if not source_path.exists():
            raise FileNotFoundError(f"Missing source file: {source_path}")

        title = row.get("title", "").strip() or source_path.stem.replace("-", " ").title()
        subject = row.get("subject", "").strip() or title
        output = row.get("output", "").strip() or f"docs/{source_path.with_suffix('.docx').name}"
        author = row.get("author", "").strip() or DEFAULT_AUTHOR
        output_path = root / output

        build_docx(source_path, title, subject, output_path, author)
        outputs.append(output_path)

    return outputs


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate Word .docx files from project Markdown sources.")
    parser.add_argument(
        "--project",
        default=".",
        help="Project folder containing docx_sources.csv. Defaults to the current folder.",
    )
    parser.add_argument(
        "--config",
        default="docx_sources.csv",
        help="Config CSV path relative to the project folder.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    outputs = generate_docs(Path(args.project), args.config)
    for output in outputs:
        print(f"Wrote {output}")


if __name__ == "__main__":
    main()
